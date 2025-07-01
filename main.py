import streamlit as st
import os
import tempfile
import subprocess
import re
import time
import json
import sys
from io import BytesIO
from datetime import datetime
from queue import Queue
import threading

from moviepy.editor import VideoFileClip
from pydub import AudioSegment
import torch
import whisper
import google.generativeai as genai
from google.generativeai.types import GenerationConfig

# ——————————————————— Streamlit configurações básicas
st.set_page_config(page_title="SIA — Secretaria de IA do Piauí", layout="wide")

# ——————————————————— Bibliotecas para exportação
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ python-docx não instalado. Instale com: pip install python-docx")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# ——————————————————— Configurações Gerais
SEGMENT_SECONDS = 15 * 60  # 15 minutos por bloco
tmp_dir = tempfile.gettempdir()

# Configuração da API Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")


# ——————————————————— Funções de Log
def log_info(message: str, show_time: bool = True):
    timestamp = datetime.now().strftime("%H:%M:%S") if show_time else ""
    formatted_msg = f"[{timestamp}] {message}" if show_time else message
    print(formatted_msg)
    sys.stdout.flush()

def log_error(message: str):
    error_msg = f"❌ ERRO [{datetime.now().strftime('%H:%M:%S')}]: {message}"
    print(error_msg)
    sys.stdout.flush()

def log_success(message: str):
    success_msg = f"✅ SUCESSO [{datetime.now().strftime('%H:%M:%S')}]: {message}"
    print(success_msg)
    sys.stdout.flush()

# ——————————————————— CSS customizado
st.markdown(
    """
    <style>
    .reportview-container .main { 
        background-color: #f0f2f6; 
    }
    h1 {
        color: #2c3e50;
    }
    .stButton>button {
        background-color: #27ae60;
        color: white;
        border-radius: 5px;
        padding: 0.4em 0.8em;
    }
    .stProgress>div>div>div>div {
        background-color: #27ae60;
    }
    textarea {
        background-color: #ffffff;
        color: #2c3e50;
        border: 1px solid #d0d3d4;
    }
    .stInfo {
        color: #34495e;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ——————————————————— Estado de Sessão
def init_session_state():
    defaults = {
        "content_txt": "",
        "content_title": "",
        "resumo": "",
        "processing": False,
        "progress_text": "",
        "progress_value": 0,
        "content_type": "",
        "current_step": "",
        "start_time": None,
    }
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)

init_session_state()

# ——————————————————— Carregamento do Whisper (cacheado)
@st.cache_resource
def load_whisper():
    log_info("🔄 Carregando modelo Whisper...")
    device = "cuda" if torch.cuda.is_available() else "cpu"
    log_info(f"🖥️ Usando dispositivo: {device}")
    model = whisper.load_model("small", device=device, in_memory=True)
    log_success("Modelo Whisper carregado com sucesso!")
    return model

whisper_model = load_whisper()

# ——————————————————— Utilitários auxiliares
sanitize = lambda s: re.sub(r'[<>:"/\\|?*]', "_", s)[:100]

def update_progress(text: str, value: int = None, step: str = ""):
    """Atualiza descrição de progresso exibida na interface"""
    st.session_state.progress_text = text
    st.session_state.current_step = step
    if value is not None:
        st.session_state.progress_value = max(0, min(100, value))
    log_info(f"📊 Progresso: {value}% - {text}")

def ensure_ffmpeg():
    log_info("🔍 Verificando FFmpeg...")
    try:
        subprocess.run(["ffprobe", "-version"], check=True, capture_output=True)
        log_success("FFmpeg encontrado e funcionando")
    except FileNotFoundError:
        log_error("FFmpeg não encontrado no PATH")
        raise RuntimeError("Instale FFmpeg (ffprobe) e deixe disponível no PATH.")

def yt_download(url: str):
    """Baixa o áudio de um vídeo do YouTube como WAV e retorna caminho + título"""
    log_info(f"🎬 Iniciando download do YouTube: {url}")
    ensure_ffmpeg()
    update_progress("📥 Baixando áudio do YouTube…", value=5, step="download")

    tmp_base = tempfile.mktemp(dir=tmp_dir)
    log_info(f"📁 Arquivo temporário: {tmp_base}")
    try:
        subprocess.run([
            "yt-dlp", "-f", "bestaudio",
            "--extract-audio", "--audio-format", "wav",
            "-o", tmp_base + ".%(ext)s", url,
        ], check=True, capture_output=True)

        wav_path = tmp_base + ".wav"
        log_success(f"Download concluído: {wav_path}")

        log_info("📋 Obtendo título do vídeo...")
        title = subprocess.check_output(
            ["yt-dlp", "--print", "%(title)s", "--no-download", url], text=True
        ).strip()

        log_success(f"Título obtido: {title}")
        return wav_path, sanitize(title)
    except subprocess.CalledProcessError as e:
        log_error(f"Erro no yt-dlp: {e}")
        raise

def video_to_wav(path: str):
    """Extrai e salva o áudio de um arquivo de vídeo em WAV temporário"""
    log_info(f"🎵 Extraindo áudio do vídeo: {path}")
    update_progress("🎵 Extraindo áudio do vídeo…", value=15, step="extract_audio")

    out_wav = tempfile.mktemp(dir=tmp_dir, suffix=".wav")
    log_info(f"💾 Salvando áudio em: {out_wav}")
    try:
        video_clip = VideoFileClip(path)
        log_info(f"⏱️ Duração do vídeo: {video_clip.duration} segundos")
        video_clip.audio.write_audiofile(out_wav, logger=None, verbose=False)
        video_clip.close()

        if os.path.exists(out_wav):
            file_size = os.path.getsize(out_wav) / (1024*1024)
            log_success(f"Áudio extraído com sucesso! Tamanho: {file_size:.2f} MB")
        else:
            log_error("Arquivo de áudio não foi criado")

        return out_wav
    except Exception as e:
        log_error(f"Erro ao extrair áudio: {e}")
        raise

def audio_to_wav(path: str):
    """Converte arquivo de áudio para WAV se necessário"""
    log_info(f"🎵 Processando arquivo de áudio: {path}")
    update_progress("🎵 Processando arquivo de áudio…", value=15, step="convert_audio")

    if path.lower().endswith('.wav'):
        log_info("✅ Arquivo já está em formato WAV")
        return path

    out_wav = tempfile.mktemp(dir=tmp_dir, suffix=".wav")
    log_info(f"🔄 Convertendo para WAV: {out_wav}")
    try:
        audio = AudioSegment.from_file(path)
        log_info(f"⏱️ Duração do áudio: {len(audio)/1000:.2f} segundos")
        audio.export(out_wav, format="wav")
        if os.path.exists(out_wav):
            file_size = os.path.getsize(out_wav) / (1024*1024)
            log_success(f"Conversão concluída! Tamanho: {file_size:.2f} MB")
        return out_wav
    except Exception as e:
        log_error(f"Erro na conversão: {e}")
        raise

def split_audio(wav_path: str):
    """Divide áudio em partes de até SEGMENT_SECONDS e retorna lista de caminhos"""
    log_info(f"✂️ Dividindo áudio em segmentos de {SEGMENT_SECONDS/60} minutos")
    update_progress("✂️ Dividindo áudio em segmentos…", value=25, step="split_audio")

    try:
        audio = AudioSegment.from_wav(wav_path)
        duration_seconds = len(audio) / 1000
        total_segments = int(duration_seconds / SEGMENT_SECONDS) + 1

        log_info(f"📊 Duração total: {duration_seconds:.2f}s - Segmentos: {total_segments}")

        parts = []
        for i in range(0, len(audio), SEGMENT_SECONDS * 1000):
            segment_num = i // (SEGMENT_SECONDS * 1000) + 1
            part_path = tempfile.mktemp(dir=tmp_dir, suffix=f"_seg{segment_num:03d}.wav")
            segment = audio[i : i + SEGMENT_SECONDS * 1000]
            segment.export(part_path, format="wav")
            parts.append(part_path)
            log_info(f"📦 Segmento {segment_num}/{total_segments} criado: {os.path.basename(part_path)}")

        log_success(f"Áudio dividido em {len(parts)} segmentos")
        return parts
    except Exception as e:
        log_error(f"Erro ao dividir áudio: {e}")
        raise

def transcribe_parts(paths):
    """Transcreve cada segmento de áudio com Whisper e concatena o resultado"""
    log_info(f"🎙️ Iniciando transcrição de {len(paths)} segmentos")
    texts = []
    total = len(paths)

    for idx, p in enumerate(paths, 1):
        pct = int((idx / total) * 70)
        update_progress(f"🎙️ Transcrevendo segmento {idx}/{total}…", value=25 + pct, step="transcribe")
        log_info(f"🔄 Processando segmento {idx}/{total}: {os.path.basename(p)}")

        try:
            if not os.path.exists(p):
                log_error(f"Arquivo não encontrado: {p}")
                continue
            file_size = os.path.getsize(p)
            if file_size == 0:
                log_info(f"⚠️ Arquivo vazio ignorado: {os.path.basename(p)}")
                continue

            log_info(f"📏 Tamanho do arquivo: {file_size / 1024:.2f} KB")
            log_info(f"🤖 Chamando Whisper para segmento {idx}...")
            start_time = time.time()
            result = whisper_model.transcribe(p)["text"].strip()
            transcribe_time = time.time() - start_time
            log_info(f"⏱️ Transcrição do segmento {idx} levou {transcribe_time:.2f}s")
            log_info(f"📝 Texto transcrito ({len(result)} chars): {result[:100]}{'...' if len(result) > 100 else ''}")
            texts.append(result)
        except Exception as e:
            log_error(f"Erro ao transcrever segmento {idx}: {e}")
        finally:
            if os.path.exists(p):
                os.remove(p)
                log_info(f"🗑️ Arquivo temporário removido: {os.path.basename(p)}")

    update_progress("🎉 Transcrição concluída!", value=95, step="complete")
    final_text = "\n".join(texts)
    log_success(f"Transcrição finalizada! Total de caracteres: {len(final_text)}")
    return final_text

# ——————————————————— Resumo com Gemini
def summarize_with_gemini(text: str) -> str:
    """Gera resumo em um único parágrafo usando o cliente genai para Gemini"""
    log_info("🤖 Iniciando geração de resumo com Gemini...")
    if not GEMINI_API_KEY:
        log_error("GEMINI_API_KEY não configurada")
        raise RuntimeError("GEMINI_API_KEY não configurada")

    model_name = "gemini-1.5-flash-latest"
    log_info(f"🔧 Usando modelo: {model_name}")
    model = genai.GenerativeModel(model_name)

    prompt = f"""
Você é um assistente que faz resumos concisos.
Leia o texto abaixo e responda com UM PARÁGRAFO ÚNICO, capturando a ideia principal e os pontos-chave sem copiar trechos literais.

Texto a ser resumido:
{text}
    """

    try:
        log_info(f"📤 Enviando {len(text)} caracteres para o Gemini...")
        start_time = time.time()
        response = model.generate_content(
            prompt,
            generation_config=GenerationConfig(
                temperature=0.2,
                top_p=0.9,
                max_output_tokens=300,
            )
        )
        api_time = time.time() - start_time
        summary = response.text.strip()
        log_info(f"⏱️ Chamada da API levou {api_time:.2f}s")
        log_success(f"Resumo gerado ({len(summary)} chars): {summary[:100]}{'...' if len(summary) > 100 else ''}")
        return summary
    except Exception as e:
        log_error(f"Erro na API do Gemini: {e}")
        raise

# ——————————————————— Funções de Exportação
def create_docx_content(title: str, transcription: str, summary: str = None):
    """Cria documento DOCX com transcrição e resumo"""
    if not DOCX_AVAILABLE:
        return None
    log_info("📄 Criando documento DOCX...")
    doc = Document()
    doc.add_heading(f'Transcrição: {title}', 0)
    if summary:
        doc.add_heading('Resumo', level=1)
        doc.add_paragraph(summary)
        doc.add_page_break()
    doc.add_heading('Transcrição Completa', level=1)
    doc.add_paragraph(transcription)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    log_success("Documento DOCX criado com sucesso")
    return buffer.getvalue()

def create_json_content(title: str, transcription: str, summary: str = None):
    """Cria conteúdo JSON estruturado"""
    log_info("📊 Criando conteúdo JSON...")
    data = {
        "titulo": title,
        "transcricao": transcription,
        "resumo": summary if summary else "",
        "data_processamento": time.strftime("%Y-%m-%d %H:%M:%S")
    }
    json_str = json.dumps(data, ensure_ascii=False, indent=2)
    log_success("Conteúdo JSON criado com sucesso")
    return json_str

def create_csv_content(title: str, transcription: str, summary: str = None):
    """Cria conteúdo CSV"""
    log_info("📈 Criando conteúdo CSV...")
    if not PANDAS_AVAILABLE:
        content = "Tipo,Conteudo\n"
        title_escaped = title.replace('"', '""')
        content += f'"Titulo","{title_escaped}"\n'
        if summary:
            summary_escaped = summary.replace('"', '""')
            content += f'"Resumo","{summary_escaped}"\n'
        transcription_escaped = transcription.replace('"', '""')
        content += f'"Transcricao","{transcription_escaped}"\n'
        log_success("CSV criado manualmente (sem pandas)")
        return content
    data = [{"Tipo": "Titulo", "Conteudo": title}]
    if summary:
        data.append({"Tipo": "Resumo", "Conteudo": summary})
    data.append({"Tipo": "Transcricao", "Conteudo": transcription})
    df = pd.DataFrame(data)
    csv_str = df.to_csv(index=False)
    log_success("CSV criado com pandas")
    return csv_str

# ——————————————————— Processamento
def process_video_sync(uploaded_file):
    """Fluxo síncrono para vídeos com logs detalhados"""
    log_info("=" * 50)
    log_info("🎬 INICIANDO PROCESSAMENTO DE VÍDEO")
    log_info(f"📁 Nome do arquivo: {uploaded_file.name}")
    log_info(f"📏 Tamanho: {uploaded_file.size / (1024*1024):.2f} MB")
    log_info("=" * 50)

    start_time = time.time()
    st.session_state.start_time = start_time

    try:
        log_info("💾 Salvando arquivo temporário...")
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tf:
            tf.write(uploaded_file.read())
            local_path = tf.name

        log_success(f"Arquivo salvo em: {local_path}")
        wav = video_to_wav(local_path)
        os.remove(local_path)
        log_info(f"🗑️ Arquivo original removido: {local_path}")

        parts = split_audio(wav)
        os.remove(wav)
        log_info(f"🗑️ Arquivo WAV principal removido: {wav}")

        transcription = transcribe_parts(parts)
        title = os.path.splitext(uploaded_file.name)[0]

        total_time = time.time() - start_time
        log_success(f"🎉 PROCESSAMENTO CONCLUÍDO EM {total_time:.2f} SEGUNDOS!")
        update_progress("✅ Transcrição completa!", value=100, step="finished")
        return "success", transcription, title

    except Exception as e:
        error_msg = f"Erro no processamento de vídeo: {str(e)}"
        log_error(error_msg)
        return "error", error_msg, ""

def process_audio_sync(uploaded_file):
    """Fluxo síncrono para arquivos de áudio com logs detalhados"""
    log_info("=" * 50)
    log_info("🎵 INICIANDO PROCESSAMENTO DE ÁUDIO")
    log_info(f"📁 Nome do arquivo: {uploaded_file.name}")
    log_info(f"📏 Tamanho: {uploaded_file.size / (1024*1024):.2f} MB")
    log_info("=" * 50)

    start_time = time.time()
    st.session_state.start_time = start_time

    try:
        log_info("💾 Salvando arquivo temporário...")
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tf:
            tf.write(uploaded_file.read())
            local_path = tf.name

        log_success(f"Arquivo salvo em: {local_path}")
        wav = audio_to_wav(local_path)
        if wav != local_path:
            os.remove(local_path)
            log_info(f"🗑️ Arquivo original removido: {local_path}")

        parts = split_audio(wav)
        os.remove(wav)
        log_info(f"🗑️ Arquivo WAV principal removido: {wav}")

        transcription = transcribe_parts(parts)
        title = os.path.splitext(uploaded_file.name)[0]

        total_time = time.time() - start_time
        log_success(f"🎉 PROCESSAMENTO CONCLUÍDO EM {total_time:.2f} SEGUNDOS!")
        update_progress("✅ Transcrição completa!", value=100, step="finished")
        return "success", transcription, title

    except Exception as e:
        error_msg = f"Erro no processamento de áudio: {str(e)}"
        log_error(error_msg)
        return "error", error_msg, ""

def process_youtube_sync(url: str):
    """Fluxo síncrono de YouTube com logs detalhados"""
    log_info("=" * 50)
    log_info("🎬 INICIANDO PROCESSAMENTO DO YOUTUBE")
    log_info(f"🔗 URL: {url}")
    log_info("=" * 50)

    start_time = time.time()
    st.session_state.start_time = start_time

    try:
        wav, title = yt_download(url)
        parts = split_audio(wav)
        os.remove(wav)
        log_info(f"🗑️ Arquivo WAV principal removido: {wav}")

        transcription = transcribe_parts(parts)
        total_time = time.time() - start_time
        log_success(f"🎉 PROCESSAMENTO CONCLUÍDO EM {total_time:.2f} SEGUNDOS!")
        update_progress("✅ Transcrição completa!", value=100, step="finished")
        return "success", transcription, title

    except Exception as e:
        error_msg = f"Erro no processamento do YouTube: {str(e)}"
        log_error(error_msg)
        return "error", error_msg, ""

# ——————————————————— Interface Streamlit
st.title("🎙️ SIA")
st.markdown(
    "<div style='font-size:14px; color:#34495e;'>Sistema de Transcrição e Resumo com IA</div>",
    unsafe_allow_html=True
)

mode = st.radio(
    "Escolha o modo:",
    ["Arquivo de Vídeo", "Arquivo de Áudio", "YouTube"],
    horizontal=True,
    key="modo_principal"
)

# ———————————————————————————
# 1) Indicador de processamento (barra + porcentagem)
# ———————————————————————————
if st.session_state.processing:
    # Exibe porcentagem geral no topo
    st.write(f"**Progresso geral:** {st.session_state.progress_value}%")
    # Informação de status + barra gráfica
    st.info(f"🔄 **PROCESSANDO** - {st.session_state.progress_text}")
    st.progress(st.session_state.progress_value)

    # Tempo decorrido e etapa atual
    if st.session_state.start_time:
        elapsed = time.time() - st.session_state.start_time
        elapsed_str = f"⏱️ Tempo decorrido: {elapsed:.0f}s"
    else:
        elapsed_str = ""

    col1, col2 = st.columns(2)
    with col1:
        st.write(f"📊 Progresso: {st.session_state.progress_value}%")
    with col2:
        if elapsed_str:
            st.write(elapsed_str)

    if st.session_state.current_step:
        st.write(f"🔧 Etapa atual: {st.session_state.current_step}")

    st.write("💡 **Dica**: Acompanhe os logs detalhados no console/terminal onde você executou o Streamlit")

# ———————————————————————————
# 2) Seção “Arquivo de Vídeo”
# ———————————————————————————
if mode == "Arquivo de Vídeo" and not st.session_state.processing:
    up_video = st.file_uploader(
        "Selecione um vídeo",
        type=["mp4", "mkv", "mov", "avi", "webm"],
        key="upload_video"
    )
    if up_video and st.button("🚀 Transcrever Vídeo", disabled=st.session_state.processing, key="btn_transcrever_video"):
        log_info("🎬 USUÁRIO CLICOU EM TRANSCREVER VÍDEO!")
        log_info(f"📋 Arquivo selecionado: {up_video.name}")

        st.session_state.processing = True
        st.session_state.progress_value = 0
        st.session_state.progress_text = "🚀 Iniciando processamento…"
        st.session_state.start_time = time.time()

        st.success("✅ Processamento iniciado! Acompanhe o progresso abaixo e os logs no console.")

        status, content, title = process_video_sync(up_video)
        if status == "success":
            st.session_state.content_txt = content
            st.session_state.content_title = title
            st.session_state.content_type = "video"
            st.session_state.processing = False
            st.success("✨ Transcrição finalizada!")
            log_success("🎉 TRANSCRIÇÃO EXIBIDA NA INTERFACE!")
        else:
            st.session_state.processing = False
            st.error(f"❌ Erro: {content}")
            log_error(f"ERRO EXIBIDO NA INTERFACE: {content}")

# ———————————————————————————
# 3) Seção “Arquivo de Áudio”
# ———————————————————————————
elif mode == "Arquivo de Áudio" and not st.session_state.processing:
    up_audio = st.file_uploader(
        "Selecione um arquivo de áudio",
        type=["mp3", "wav", "m4a", "aac", "ogg", "flac", "wma"],
        key="upload_audio"
    )
    if up_audio and st.button("🚀 Transcrever Áudio", disabled=st.session_state.processing, key="btn_transcrever_audio"):
        log_info("🎵 USUÁRIO CLICOU EM TRANSCREVER ÁUDIO!")
        log_info(f"📋 Arquivo selecionado: {up_audio.name}")

        st.session_state.processing = True
        st.session_state.progress_value = 0
        st.session_state.progress_text = "🚀 Iniciando processamento…"
        st.session_state.start_time = time.time()

        st.success("✅ Processamento iniciado! Acompanhe o progresso abaixo e os logs no console.")

        status, content, title = process_audio_sync(up_audio)
        if status == "success":
            st.session_state.content_txt = content
            st.session_state.content_title = title
            st.session_state.content_type = "audio"
            st.session_state.processing = False
            st.success("✨ Transcrição finalizada!")
            log_success("🎉 TRANSCRIÇÃO EXIBIDA NA INTERFACE!")
        else:
            st.session_state.processing = False
            st.error(f"❌ Erro: {content}")
            log_error(f"ERRO EXIBIDO NA INTERFACE: {content}")

# ———————————————————————————
# 4) Seção “YouTube”
# ———————————————————————————
elif mode == "YouTube" and not st.session_state.processing:
    url = st.text_input(
        "🔗 Cole a URL do YouTube aqui",
        key="input_youtube"
    )
    if url and st.button("🚀 Transcrever YouTube", disabled=st.session_state.processing, key="btn_transcrever_youtube"):
        log_info("🎬 USUÁRIO CLICOU EM TRANSCREVER YOUTUBE!")
        log_info(f"📋 URL inserida: {url}")

        st.session_state.processing = True
        st.session_state.progress_value = 0
        st.session_state.progress_text = "🚀 Iniciando download…"
        st.session_state.start_time = time.time()

        st.success("✅ Processamento iniciado! Acompanhe o progresso abaixo e os logs no console.")

        status, content, title = process_youtube_sync(url)
        if status == "success":
            st.session_state.content_txt = content
            st.session_state.content_title = title
            st.session_state.content_type = "youtube"
            st.session_state.processing = False
            st.success("✨ Transcrição finalizada!")
            log_success("🎉 TRANSCRIÇÃO EXIBIDA NA INTERFACE!")
        else:
            st.session_state.processing = False
            st.error(f"❌ Erro: {content}")
            log_error(f"ERRO EXIBIDO NA INTERFACE: {content}")

# ———————————————————————————
# 5) Exibição e geração de resumo
# ———————————————————————————
if st.session_state.content_txt:
    col1, col2 = st.columns([3, 1])
    with col1:
        with st.expander("📄 Ver Transcrição Completa", expanded=False):
            st.text_area("", st.session_state.content_txt, height=300)

    with col2:
        if st.button("🤖 Gerar Resumo", key="btn_gerar_resumo", disabled=st.session_state.processing):
            with st.spinner("Gerando resumo com Gemini…"):
                try:
                    resumo_text = summarize_with_gemini(st.session_state.content_txt)
                    st.session_state.resumo = resumo_text
                    st.success("✅ Resumo gerado!")
                except Exception as e:
                    st.error(f"Erro ao gerar resumo: {e}")

# ———————————————————————————
# 6) Exibição do Resumo
# ———————————————————————————
if st.session_state.resumo:
    st.subheader("📋 Resumo")
    st.info(st.session_state.resumo)

# ———————————————————————————
# 7) Seção de exportação
# ———————————————————————————
if st.session_state.content_txt:
    st.subheader("📥 Exportar Conteúdo")
    col1, col2, col3 = st.columns(3)

    export_transcription = st.checkbox(
        "Incluir Transcrição",
        value=True,
        key="chk_transcricao",
        disabled=st.session_state.processing
    )
    export_summary = st.checkbox(
        "Incluir Resumo",
        value=bool(st.session_state.resumo),
        disabled=not st.session_state.resumo or st.session_state.processing,
        key="chk_resumo"
    )

    with col1:
        if export_transcription or export_summary:
            content_to_export = ""
            if export_summary and st.session_state.resumo:
                content_to_export += f"RESUMO:\n{st.session_state.resumo}\n\n"
            if export_transcription:
                content_to_export += f"TRANSCRIÇÃO COMPLETA:\n{st.session_state.content_txt}"
            st.download_button(
                label="📄 Baixar como TXT",
                data=content_to_export,
                file_name=f"{sanitize(st.session_state.content_title)}.txt",
                mime="text/plain",
                key="dl_txt"
            )

    with col2:
        if export_transcription or export_summary:
            json_content = create_json_content(
                st.session_state.content_title,
                st.session_state.content_txt if export_transcription else "",
                st.session_state.resumo if export_summary else None
            )
            st.download_button(
                label="📊 Baixar como JSON",
                data=json_content,
                file_name=f"{sanitize(st.session_state.content_title)}.json",
                mime="application/json",
                key="dl_json"
            )

    with col3:
        if export_transcription or export_summary:
            csv_content = create_csv_content(
                st.session_state.content_title,
                st.session_state.content_txt if export_transcription else "",
                st.session_state.resumo if export_summary else None
            )
            st.download_button(
                label="📈 Baixar como CSV",
                data=csv_content,
                file_name=f"{sanitize(st.session_state.content_title)}.csv",
                mime="text/csv",
                key="dl_csv"
            )

    if DOCX_AVAILABLE and (export_transcription or export_summary):
        docx_content = create_docx_content(
            st.session_state.content_title,
            st.session_state.content_txt if export_transcription else "",
            st.session_state.resumo if export_summary else None
        )
        if docx_content:
            st.download_button(
                label="📝 Baixar como DOCX (Word)",
                data=docx_content,
                file_name=f"{sanitize(st.session_state.content_title)}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_docx"
            )

# Rodapé


st.markdown("---")
st.markdown("🏛️ **SIA — Secretaria de IA do Piauí** | Sistema completo de transcrição e análise de conteúdo audiovisual")
